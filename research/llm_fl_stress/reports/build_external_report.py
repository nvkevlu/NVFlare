# Copyright (c) 2026, NVIDIA CORPORATION.  All rights reserved.
#
# Licensed under the Apache License, Version 2.0 (the "License");
# you may not use this file except in compliance with the License.
# You may obtain a copy of the License at
#
#     http://www.apache.org/licenses/LICENSE-2.0
#
# Unless required by applicable law or agreed to in writing, software
# distributed under the License is distributed on an "AS IS" BASIS,
# WITHOUT WARRANTIES OR CONDITIONS OF ANY KIND, either express or implied.
# See the License for the specific language governing permissions and
# limitations under the License.

from __future__ import annotations

import sys
from pathlib import Path

from PIL import Image, ImageDraw, ImageFont
from docx import Document
from docx.enum.section import WD_SECTION_START
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor, Twips


ROOT = Path(__file__).resolve().parents[3]
REPORT_DIR = Path(__file__).resolve().parent
ASSET_DIR = REPORT_DIR / "assets"
OUTPUT_PATH = REPORT_DIR / "NVFLARE_Large_Model_Stress_Test_Report_2026-07-21.docx"

SKILL_DIR = Path(
    "/Users/kevlu/.codex/plugins/cache/openai-primary-runtime/documents/26.715.12143/skills/documents"
)
sys.path.insert(0, str(SKILL_DIR / "scripts"))
from table_geometry import apply_table_geometry  # noqa: E402


PAGE_WIDTH_DXA = 9360
TABLE_INDENT_DXA = 120
CELL_MARGINS = {"top": 100, "bottom": 100, "start": 120, "end": 120}

NAVY = "17365D"
BLUE = "2E74B5"
DARK_BLUE = "1F4D78"
TEAL = "2A7F8E"
LIGHT_BLUE = "E8EEF5"
LIGHT_TEAL = "E7F2F4"
LIGHT_GRAY = "F2F4F7"
MID_GRAY = "667085"
DARK_GRAY = "344054"
GOLD = "B07D18"
LIGHT_GOLD = "FFF4D6"
RED = "9B1C1C"
LIGHT_RED = "FCE8E6"
GREEN = "276749"
LIGHT_GREEN = "E7F4EC"
WHITE = "FFFFFF"
BLACK = "1D2939"


def rgb(hex_color: str) -> RGBColor:
    return RGBColor.from_string(hex_color)


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_border(cell, **edges) -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    borders = tc_pr.first_child_found_in("w:tcBorders")
    if borders is None:
        borders = OxmlElement("w:tcBorders")
        tc_pr.append(borders)
    for edge_name, edge_data in edges.items():
        edge = borders.find(qn(f"w:{edge_name}"))
        if edge is None:
            edge = OxmlElement(f"w:{edge_name}")
            borders.append(edge)
        for key, value in edge_data.items():
            edge.set(qn(f"w:{key}"), str(value))


def set_paragraph_shading(paragraph, fill: str) -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    shd = p_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        p_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_paragraph_left_border(paragraph, color: str, size: int = 18, space: int = 8) -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    p_bdr = p_pr.find(qn("w:pBdr"))
    if p_bdr is None:
        p_bdr = OxmlElement("w:pBdr")
        p_pr.append(p_bdr)
    left = p_bdr.find(qn("w:left"))
    if left is None:
        left = OxmlElement("w:left")
        p_bdr.append(left)
    left.set(qn("w:val"), "single")
    left.set(qn("w:sz"), str(size))
    left.set(qn("w:space"), str(space))
    left.set(qn("w:color"), color)


def set_paragraph_bottom_border(paragraph, color: str, size: int = 10, space: int = 8) -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    p_bdr = p_pr.find(qn("w:pBdr"))
    if p_bdr is None:
        p_bdr = OxmlElement("w:pBdr")
        p_pr.append(p_bdr)
    bottom = p_bdr.find(qn("w:bottom"))
    if bottom is None:
        bottom = OxmlElement("w:bottom")
        p_bdr.append(bottom)
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), str(size))
    bottom.set(qn("w:space"), str(space))
    bottom.set(qn("w:color"), color)


def set_repeat_table_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def set_keep_with_next(paragraph, value: bool = True) -> None:
    paragraph.paragraph_format.keep_with_next = value


def set_run_font(run, name: str = "Calibri", size: float | None = None, color: str | None = None,
                 bold: bool | None = None, italic: bool | None = None) -> None:
    run.font.name = name
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), name)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), name)
    if size is not None:
        run.font.size = Pt(size)
    if color is not None:
        run.font.color.rgb = rgb(color)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def configure_styles(document: Document) -> None:
    styles = document.styles

    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    normal.font.size = Pt(11)
    normal.font.color.rgb = rgb(BLACK)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.10

    for style_name, size, color, before, after in (
        ("Heading 1", 16, BLUE, 16, 8),
        ("Heading 2", 13, BLUE, 12, 6),
        ("Heading 3", 12, DARK_BLUE, 8, 4),
    ):
        style = styles[style_name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = rgb(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    caption = styles["Caption"]
    caption.font.name = "Calibri"
    caption._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    caption._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    caption.font.size = Pt(9)
    caption.font.italic = True
    caption.font.color.rgb = rgb(MID_GRAY)
    caption.paragraph_format.space_before = Pt(4)
    caption.paragraph_format.space_after = Pt(9)
    caption.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER

    for style_name in ("Header", "Footer"):
        style = styles[style_name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
        style.font.size = Pt(8.5)
        style.font.color.rgb = rgb(MID_GRAY)


def add_custom_numbering(document: Document) -> tuple[int, int]:
    numbering = document.part.numbering_part.element
    nsid = max([int(n.get(qn("w:abstractNumId"))) for n in numbering.findall(qn("w:abstractNum"))] or [0]) + 1
    bullet_abstract = OxmlElement("w:abstractNum")
    bullet_abstract.set(qn("w:abstractNumId"), str(nsid))
    multi = OxmlElement("w:multiLevelType")
    multi.set(qn("w:val"), "singleLevel")
    bullet_abstract.append(multi)
    level = OxmlElement("w:lvl")
    level.set(qn("w:ilvl"), "0")
    start = OxmlElement("w:start")
    start.set(qn("w:val"), "1")
    num_fmt = OxmlElement("w:numFmt")
    num_fmt.set(qn("w:val"), "bullet")
    lvl_text = OxmlElement("w:lvlText")
    lvl_text.set(qn("w:val"), "•")
    lvl_jc = OxmlElement("w:lvlJc")
    lvl_jc.set(qn("w:val"), "left")
    p_pr = OxmlElement("w:pPr")
    tabs = OxmlElement("w:tabs")
    tab = OxmlElement("w:tab")
    tab.set(qn("w:val"), "num")
    tab.set(qn("w:pos"), "720")
    tabs.append(tab)
    indent = OxmlElement("w:ind")
    indent.set(qn("w:left"), "720")
    indent.set(qn("w:hanging"), "360")
    spacing = OxmlElement("w:spacing")
    spacing.set(qn("w:after"), "160")
    spacing.set(qn("w:line"), "280")
    spacing.set(qn("w:lineRule"), "auto")
    p_pr.extend([tabs, indent, spacing])
    r_pr = OxmlElement("w:rPr")
    fonts = OxmlElement("w:rFonts")
    fonts.set(qn("w:ascii"), "Calibri")
    fonts.set(qn("w:hAnsi"), "Calibri")
    r_pr.append(fonts)
    level.extend([start, num_fmt, lvl_text, lvl_jc, p_pr, r_pr])
    bullet_abstract.append(level)
    numbering.append(bullet_abstract)

    bullet_num_id = max([int(n.get(qn("w:numId"))) for n in numbering.findall(qn("w:num"))] or [0]) + 1
    bullet_num = OxmlElement("w:num")
    bullet_num.set(qn("w:numId"), str(bullet_num_id))
    abstract_id = OxmlElement("w:abstractNumId")
    abstract_id.set(qn("w:val"), str(nsid))
    bullet_num.append(abstract_id)
    numbering.append(bullet_num)

    decimal_abstract_id = nsid + 1
    decimal_abstract = OxmlElement("w:abstractNum")
    decimal_abstract.set(qn("w:abstractNumId"), str(decimal_abstract_id))
    decimal_multi = OxmlElement("w:multiLevelType")
    decimal_multi.set(qn("w:val"), "singleLevel")
    decimal_abstract.append(decimal_multi)
    decimal_level = OxmlElement("w:lvl")
    decimal_level.set(qn("w:ilvl"), "0")
    decimal_start = OxmlElement("w:start")
    decimal_start.set(qn("w:val"), "1")
    decimal_fmt = OxmlElement("w:numFmt")
    decimal_fmt.set(qn("w:val"), "decimal")
    decimal_text = OxmlElement("w:lvlText")
    decimal_text.set(qn("w:val"), "%1.")
    decimal_jc = OxmlElement("w:lvlJc")
    decimal_jc.set(qn("w:val"), "left")
    decimal_p_pr = OxmlElement("w:pPr")
    decimal_tabs = OxmlElement("w:tabs")
    decimal_tab = OxmlElement("w:tab")
    decimal_tab.set(qn("w:val"), "num")
    decimal_tab.set(qn("w:pos"), "720")
    decimal_tabs.append(decimal_tab)
    decimal_indent = OxmlElement("w:ind")
    decimal_indent.set(qn("w:left"), "720")
    decimal_indent.set(qn("w:hanging"), "360")
    decimal_spacing = OxmlElement("w:spacing")
    decimal_spacing.set(qn("w:after"), "160")
    decimal_spacing.set(qn("w:line"), "280")
    decimal_spacing.set(qn("w:lineRule"), "auto")
    decimal_p_pr.extend([decimal_tabs, decimal_indent, decimal_spacing])
    decimal_level.extend([decimal_start, decimal_fmt, decimal_text, decimal_jc, decimal_p_pr])
    decimal_abstract.append(decimal_level)
    numbering.append(decimal_abstract)

    decimal_num_id = bullet_num_id + 1
    decimal_num = OxmlElement("w:num")
    decimal_num.set(qn("w:numId"), str(decimal_num_id))
    decimal_ref = OxmlElement("w:abstractNumId")
    decimal_ref.set(qn("w:val"), str(decimal_abstract_id))
    decimal_num.append(decimal_ref)
    numbering.append(decimal_num)
    return bullet_num_id, decimal_num_id


def apply_numbering(paragraph, num_id: int) -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    num_pr = p_pr.find(qn("w:numPr"))
    if num_pr is None:
        num_pr = OxmlElement("w:numPr")
        p_pr.append(num_pr)
    ilvl = OxmlElement("w:ilvl")
    ilvl.set(qn("w:val"), "0")
    num_id_element = OxmlElement("w:numId")
    num_id_element.set(qn("w:val"), str(num_id))
    num_pr.extend([ilvl, num_id_element])


def add_bullet(document: Document, text: str, bullet_num_id: int, *, bold_prefix: str | None = None) -> None:
    paragraph = document.add_paragraph()
    apply_numbering(paragraph, bullet_num_id)
    if bold_prefix and text.startswith(bold_prefix):
        run = paragraph.add_run(bold_prefix)
        set_run_font(run, bold=True)
        remainder = paragraph.add_run(text[len(bold_prefix):])
        set_run_font(remainder)
    else:
        run = paragraph.add_run(text)
        set_run_font(run)


def add_numbered(document: Document, text: str, decimal_num_id: int) -> None:
    paragraph = document.add_paragraph()
    apply_numbering(paragraph, decimal_num_id)
    run = paragraph.add_run(text)
    set_run_font(run)


def add_callout(document: Document, label: str, text: str, *, fill: str = LIGHT_BLUE, color: str = BLUE) -> None:
    paragraph = document.add_paragraph()
    paragraph.paragraph_format.left_indent = Inches(0.12)
    paragraph.paragraph_format.right_indent = Inches(0.08)
    paragraph.paragraph_format.space_before = Pt(5)
    paragraph.paragraph_format.space_after = Pt(10)
    set_paragraph_shading(paragraph, fill)
    set_paragraph_left_border(paragraph, color)
    label_run = paragraph.add_run(f"{label}  ")
    set_run_font(label_run, bold=True, color=color)
    text_run = paragraph.add_run(text)
    set_run_font(text_run, color=BLACK)


def add_caption(document: Document, text: str) -> None:
    paragraph = document.add_paragraph(style="Caption")
    paragraph.add_run(text)


def add_figure(document: Document, path: Path, width_inches: float, alt_text: str, caption: str) -> None:
    paragraph = document.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.space_before = Pt(4)
    paragraph.paragraph_format.space_after = Pt(0)
    run = paragraph.add_run()
    shape = run.add_picture(str(path), width=Inches(width_inches))
    shape._inline.docPr.set("descr", alt_text)
    add_caption(document, caption)


def add_table(
    document: Document,
    headers: list[str],
    rows: list[list[str]],
    widths_dxa: list[int],
    *,
    alignments: list[str] | None = None,
    font_size: float = 9.2,
    status_column: int | None = None,
):
    table = document.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    table.allow_autofit = False
    header_row = table.rows[0]
    set_repeat_table_header(header_row)
    for column_index, header in enumerate(headers):
        cell = header_row.cells[column_index]
        set_cell_shading(cell, LIGHT_GRAY)
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        paragraph = cell.paragraphs[0]
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        paragraph.paragraph_format.space_before = Pt(0)
        paragraph.paragraph_format.space_after = Pt(0)
        paragraph.paragraph_format.line_spacing = 1.05
        run = paragraph.add_run(header)
        set_run_font(run, size=9, bold=True, color=NAVY)

    for row_data in rows:
        row = table.add_row()
        for column_index, value in enumerate(row_data):
            cell = row.cells[column_index]
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            paragraph = cell.paragraphs[0]
            alignment = alignments[column_index] if alignments else "left"
            paragraph.alignment = {
                "left": WD_ALIGN_PARAGRAPH.LEFT,
                "center": WD_ALIGN_PARAGRAPH.CENTER,
                "right": WD_ALIGN_PARAGRAPH.RIGHT,
            }[alignment]
            paragraph.paragraph_format.space_before = Pt(0)
            paragraph.paragraph_format.space_after = Pt(0)
            paragraph.paragraph_format.line_spacing = 1.05
            run = paragraph.add_run(str(value))
            run_color = BLACK
            if status_column is not None and column_index == status_column:
                lowered = str(value).lower()
                if "pass" in lowered or "complete" in lowered:
                    run_color = GREEN
                if "warning" in lowered or "pressure" in lowered:
                    run_color = GOLD
                if "oom" in lowered or "stall" in lowered:
                    run_color = RED
                run.bold = True
            set_run_font(run, size=font_size, color=run_color)

    apply_table_geometry(
        table,
        widths_dxa,
        table_width_dxa=PAGE_WIDTH_DXA,
        indent_dxa=TABLE_INDENT_DXA,
        cell_margins_dxa=CELL_MARGINS,
    )
    document.add_paragraph().paragraph_format.space_after = Pt(1)
    return table


def add_heading(document: Document, text: str, level: int = 1) -> None:
    paragraph = document.add_paragraph(text, style=f"Heading {level}")
    set_keep_with_next(paragraph)


def add_body(document: Document, text: str, *, bold_lead: str | None = None) -> None:
    paragraph = document.add_paragraph()
    if bold_lead and text.startswith(bold_lead):
        lead = paragraph.add_run(bold_lead)
        set_run_font(lead, bold=True)
        rest = paragraph.add_run(text[len(bold_lead):])
        set_run_font(rest)
    else:
        run = paragraph.add_run(text)
        set_run_font(run)


def setup_page(section) -> None:
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)


FONT_REGULAR = Path("/System/Library/Fonts/Supplemental/Arial.ttf")
FONT_BOLD = Path("/System/Library/Fonts/Supplemental/Arial Bold.ttf")
FONT_ITALIC = Path("/System/Library/Fonts/Supplemental/Arial Italic.ttf")


def chart_font(size: int, *, bold: bool = False, italic: bool = False):
    selected = FONT_BOLD if bold else FONT_ITALIC if italic else FONT_REGULAR
    if selected.exists():
        return ImageFont.truetype(str(selected), size=size)
    return ImageFont.load_default()


def chart_color(hex_color: str):
    return tuple(int(hex_color[index : index + 2], 16) for index in (0, 2, 4))


def draw_centered_multiline(draw, box, text: str, font, fill, spacing: int = 8) -> None:
    left, top, right, bottom = box
    bbox = draw.multiline_textbbox((0, 0), text, font=font, align="center", spacing=spacing)
    width = bbox[2] - bbox[0]
    height = bbox[3] - bbox[1]
    draw.multiline_text(
        ((left + right - width) / 2, (top + bottom - height) / 2),
        text,
        font=font,
        fill=fill,
        align="center",
        spacing=spacing,
    )


def create_workflow_chart(path: Path) -> None:
    image = Image.new("RGB", (2100, 500), "white")
    draw = ImageDraw.Draw(image)
    steps = [
        (40, "1", "Server creates\nexact-size tensors", LIGHT_BLUE, NAVY),
        (450, "2", "Server sends one\nglobal model", LIGHT_BLUE, NAVY),
        (860, "3", "Clients verify and\nchange one value", LIGHT_TEAL, TEAL),
        (1270, "4", "Clients return the\nfull model state", LIGHT_TEAL, TEAL),
        (1680, "5", "Server aggregates\nand checkpoints", LIGHT_GOLD, GOLD),
    ]
    for index, (left, number, label, fill, edge) in enumerate(steps):
        box = (left, 105, left + 350, 350)
        draw.rounded_rectangle(box, radius=20, fill=chart_color(fill), outline=chart_color(edge), width=4)
        draw.text((left + 22, 125), number, font=chart_font(27, bold=True), fill=chart_color(edge))
        draw_centered_multiline(draw, box, label, chart_font(29), chart_color(BLACK), spacing=10)
        if index < len(steps) - 1:
            start_x = left + 354
            end_x = steps[index + 1][0] - 12
            y = 228
            draw.line((start_x, y, end_x, y), fill=chart_color("98A2B3"), width=4)
            draw.polygon([(end_x, y), (end_x - 18, y - 11), (end_x - 18, y + 11)], fill=chart_color("98A2B3"))
    note = "The client step is intentionally not LLM training."
    note_font = chart_font(24, italic=True)
    note_box = draw.textbbox((0, 0), note, font=note_font)
    draw.text(((2100 - (note_box[2] - note_box[0])) / 2, 425), note, font=note_font, fill=chart_color(MID_GRAY))
    image.save(path, dpi=(180, 180))


def create_scale_chart(path: Path) -> None:
    width, height = 1700, 900
    image = Image.new("RGB", (width, height), "white")
    draw = ImageDraw.Draw(image)
    left, right, top, bottom = 155, 1630, 145, 730
    models = [10.0, 12.0, 14.7, 32.5]
    series = [
        ("Server process peak", [57.38, 68.65, 83.94, 184.51], BLUE, False),
        ("Largest client process peak", [22.00, 25.91, 31.19, 66.44], TEAL, False),
        ("Raw model payload", [18.63, 22.35, 27.38, 60.54], GOLD, True),
    ]

    draw.text((left, 42), "Measured one-round, two-client memory curve", font=chart_font(34, bold=True), fill=chart_color(NAVY))
    for tick in range(0, 201, 40):
        y = bottom - (tick / 205) * (bottom - top)
        draw.line((left, y, right, y), fill=chart_color("D0D5DD"), width=2)
        label = str(tick)
        bbox = draw.textbbox((0, 0), label, font=chart_font(22))
        draw.text((left - 22 - (bbox[2] - bbox[0]), y - 12), label, font=chart_font(22), fill=chart_color(MID_GRAY))
    draw.line((left, top, left, bottom), fill=chart_color(DARK_GRAY), width=3)
    draw.line((left, bottom, right, bottom), fill=chart_color(DARK_GRAY), width=3)

    def x_for(value):
        return left + ((value - 9.0) / (34.0 - 9.0)) * (right - left)

    def y_for(value):
        return bottom - (value / 205.0) * (bottom - top)

    for label, values, color, dashed in series:
        points = [(x_for(model), y_for(value)) for model, value in zip(models, values)]
        for point_index in range(len(points) - 1):
            start = points[point_index]
            end = points[point_index + 1]
            if dashed:
                segments = 20
                for segment in range(0, segments, 2):
                    t1, t2 = segment / segments, min((segment + 1) / segments, 1)
                    draw.line(
                        (
                            start[0] + (end[0] - start[0]) * t1,
                            start[1] + (end[1] - start[1]) * t1,
                            start[0] + (end[0] - start[0]) * t2,
                            start[1] + (end[1] - start[1]) * t2,
                        ),
                        fill=chart_color(color),
                        width=5,
                    )
            else:
                draw.line((*start, *end), fill=chart_color(color), width=6)
        for x, y in points:
            draw.ellipse((x - 9, y - 9, x + 9, y + 9), fill=chart_color(color), outline="white", width=3)
        if label == "Server process peak":
            for (x, y), value in zip(points, values):
                text = f"{value:.1f}"
                bbox = draw.textbbox((0, 0), text, font=chart_font(21, bold=True))
                draw.text((x - (bbox[2] - bbox[0]) / 2, y - 40), text, font=chart_font(21, bold=True), fill=chart_color(DARK_BLUE))

    for model in models:
        x = x_for(model)
        label = f"{model:g}B"
        bbox = draw.textbbox((0, 0), label, font=chart_font(23))
        draw.text((x - (bbox[2] - bbox[0]) / 2, bottom + 18), label, font=chart_font(23), fill=chart_color(DARK_GRAY))
    axis_label = "Model parameter count"
    bbox = draw.textbbox((0, 0), axis_label, font=chart_font(23))
    draw.text(((left + right - (bbox[2] - bbox[0])) / 2, bottom + 68), axis_label, font=chart_font(23), fill=chart_color(DARK_GRAY))
    draw.text((22, 87), "Memory (GiB)", font=chart_font(23), fill=chart_color(DARK_GRAY))

    legend_x = left
    for label, _values, color, dashed in series:
        draw.line((legend_x, 98, legend_x + 70, 98), fill=chart_color(color), width=5)
        if not dashed:
            draw.ellipse((legend_x + 28, 89, legend_x + 46, 107), fill=chart_color(color), outline="white", width=2)
        draw.text((legend_x + 82, 83), label, font=chart_font(20), fill=chart_color(DARK_GRAY))
        legend_x += 480
    image.save(path, dpi=(180, 180))


def create_memory_chart(path: Path) -> None:
    width, height = 1750, 700
    image = Image.new("RGB", (width, height), "white")
    draw = ImageDraw.Draw(image)
    draw.text((70, 38), "32.5B, three clients, final-round peak: 493.91 GiB", font=chart_font(34, bold=True), fill=chart_color(NAVY))
    draw.text((70, 100), "Process-tree RSS at this instant: 225.16 GiB (overlaps anonymous memory; do not add it)", font=chart_font(22), fill=chart_color(MID_GRAY))
    draw.text((70, 136), "Linux MemAvailable: 268.20 GiB because much of the file cache could be reclaimed", font=chart_font(22), fill=chart_color(MID_GRAY))
    bar_left, bar_right, bar_top, bar_bottom = 75, 1675, 230, 385
    values = [("Anonymous / active", 224.70, BLUE), ("File cache / reclaimable", 261.57, TEAL), ("Kernel", 7.63, GOLD)]
    total = sum(value for _label, value, _color in values)
    current = bar_left
    for label, value, color in values:
        segment = (value / total) * (bar_right - bar_left)
        draw.rectangle((current, bar_top, current + segment, bar_bottom), fill=chart_color(color))
        if segment > 100:
            text = f"{value:.1f} GiB"
            bbox = draw.textbbox((0, 0), text, font=chart_font(27, bold=True))
            draw.text((current + segment / 2 - (bbox[2] - bbox[0]) / 2, 285), text, font=chart_font(27, bold=True), fill="white")
        current += segment
    draw.rectangle((bar_left, bar_top, bar_right, bar_bottom), outline=chart_color(DARK_GRAY), width=2)
    legend_x = 165
    for label, _value, color in values:
        draw.rounded_rectangle((legend_x, 455, legend_x + 34, 489), radius=5, fill=chart_color(color))
        draw.text((legend_x + 50, 454), label, font=chart_font(23), fill=chart_color(DARK_GRAY))
        legend_x += 505
    draw.text((190, 555), "Active application memory", font=chart_font(23, bold=True), fill=chart_color(BLUE))
    draw.text((670, 555), "+", font=chart_font(23, bold=True), fill=chart_color(MID_GRAY))
    draw.text((725, 555), "reclaimable operating cache", font=chart_font(23, bold=True), fill=chart_color(TEAL))
    draw.text((1235, 555), "+ kernel", font=chart_font(23, bold=True), fill=chart_color(GOLD))
    image.save(path, dpi=(180, 180))


def create_client_chart(path: Path) -> None:
    width, height = 1450, 720
    image = Image.new("RGB", (width, height), "white")
    draw = ImageDraw.Draw(image)
    draw.text((70, 35), "More clients did not create one full server copy per client", font=chart_font(31, bold=True), fill=chart_color(NAVY))
    left, right, top, bottom = 130, 1365, 125, 580
    draw.line((left, top, left, bottom), fill=chart_color(DARK_GRAY), width=3)
    draw.line((left, bottom, right, bottom), fill=chart_color(DARK_GRAY), width=3)
    for tick in range(0, 201, 50):
        y = bottom - (tick / 205) * (bottom - top)
        draw.line((left, y, right, y), fill=chart_color("D0D5DD"), width=2)
        label = str(tick)
        bbox = draw.textbbox((0, 0), label, font=chart_font(21))
        draw.text((left - 20 - (bbox[2] - bbox[0]), y - 11), label, font=chart_font(21), fill=chart_color(MID_GRAY))
    values = [180.33, 184.51, 185.09]
    labels = ["1 client", "2 clients", "3 clients"]
    fills = [LIGHT_BLUE, BLUE, TEAL]
    centers = [370, 750, 1130]
    for center, value, label, fill in zip(centers, values, labels, fills):
        bar_height = (value / 205) * (bottom - top)
        x1, x2 = center - 105, center + 105
        y1 = bottom - bar_height
        draw.rectangle((x1, y1, x2, bottom), fill=chart_color(fill), outline=chart_color(DARK_BLUE), width=3)
        value_text = f"{value:.2f} GiB"
        bbox = draw.textbbox((0, 0), value_text, font=chart_font(24, bold=True))
        draw.text((center - (bbox[2] - bbox[0]) / 2, y1 - 38), value_text, font=chart_font(24, bold=True), fill=chart_color(DARK_BLUE))
        bbox = draw.textbbox((0, 0), label, font=chart_font(23))
        draw.text((center - (bbox[2] - bbox[0]) / 2, bottom + 20), label, font=chart_font(23), fill=chart_color(DARK_GRAY))
    draw.text((20, 86), "First-round server RSS (GiB)", font=chart_font(21), fill=chart_color(DARK_GRAY))
    image.save(path, dpi=(180, 180))


def build_report() -> Path:
    ASSET_DIR.mkdir(parents=True, exist_ok=True)
    workflow_path = ASSET_DIR / "external-report-workflow.png"
    scale_path = ASSET_DIR / "external-report-scale-memory.png"
    memory_path = ASSET_DIR / "external-report-memory-breakdown.png"
    client_path = ASSET_DIR / "external-report-client-count.png"
    create_workflow_chart(workflow_path)
    create_scale_chart(scale_path)
    create_memory_chart(memory_path)
    create_client_chart(client_path)

    document = Document()
    document.settings.odd_and_even_pages_header_footer = False
    configure_styles(document)
    bullet_num_id, decimal_num_id = add_custom_numbering(document)
    setup_page(document.sections[0])

    # Cover page: editorial-cover override applied consistently to title block.
    for _ in range(5):
        document.add_paragraph().paragraph_format.space_after = Pt(8)
    kicker = document.add_paragraph()
    kicker.alignment = WD_ALIGN_PARAGRAPH.CENTER
    kicker.paragraph_format.space_after = Pt(16)
    run = kicker.add_run("TECHNICAL FINDINGS REPORT")
    set_run_font(run, size=10, bold=True, color=GOLD)

    title = document.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.space_after = Pt(10)
    run = title.add_run("NVFLARE Large-Model\nServer Stress Testing")
    set_run_font(run, size=28, bold=True, color=NAVY)

    subtitle = document.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.paragraph_format.space_after = Pt(22)
    run = subtitle.add_run("Measured results through 32.5 billion parameters, three clients, and three federated rounds")
    set_run_font(run, size=14, color=DARK_BLUE)

    rule = document.add_paragraph()
    rule.alignment = WD_ALIGN_PARAGRAPH.CENTER
    rule.paragraph_format.space_after = Pt(22)
    set_paragraph_bottom_border(rule, GOLD, size=12, space=0)

    meta = document.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    meta.paragraph_format.space_after = Pt(5)
    run = meta.add_run("Prepared for external technical review")
    set_run_font(run, size=11, bold=True, color=DARK_GRAY)
    meta = document.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = meta.add_run("Interim report • July 21, 2026")
    set_run_font(run, size=10, color=MID_GRAY)

    for _ in range(4):
        document.add_paragraph().paragraph_format.space_after = Pt(8)
    scope = document.add_paragraph()
    scope.alignment = WD_ALIGN_PARAGRAPH.CENTER
    scope.paragraph_format.left_indent = Inches(0.65)
    scope.paragraph_format.right_indent = Inches(0.65)
    run = scope.add_run(
        "Scope: CPU-based validation of full-size model distribution, return transfer, aggregation, disk offload, persistence, telemetry, and memory-failure behavior."
    )
    set_run_font(run, size=10.5, color=MID_GRAY, italic=True)

    body_section = document.add_section(WD_SECTION_START.NEW_PAGE)
    setup_page(body_section)
    body_section.different_first_page_header_footer = False

    add_heading(document, "Executive Summary", 1)
    add_body(
        document,
        "This project tested whether NVFLARE's server data path can move, receive, aggregate, and persist LLM-sized model states without requiring the cost and variability of full LLM training. The largest completed workload represented a 32.5-billion-parameter BF16 model, used three clients, and completed three federated rounds."
    )
    add_callout(
        document,
        "Bottom line",
        "NVFLARE successfully completed the 32.5B three-client exchange and aggregation workflow. The server's active process working set peaked near 245 GiB. A much larger 494 GiB cgroup reading was mostly caused by reclaimable file cache, not by a separate full model copy for each client.",
        fill=LIGHT_GREEN,
        color=GREEN,
    )
    for text in (
        "The distributed qualification curve passed from 135 million parameters through one 32.5B round with two clients.",
        "A 32.5B workload also completed two rounds with three clients and completed all three rounds with three clients; the three-round report raised a conservative memory-growth warning even though the final round-to-round change had already plateaued.",
        "The earlier two-client OOM was not a whole-machine natural OOM and was not caused by one model copy per client. It was intentionally triggered after a 235 GiB cgroup-constrained run had already stalled near 226 GiB of active anonymous memory.",
        "The available network delivered approximately 23.6 Gbps in iperf3, while NVFLARE moved the model at only about 0.75-0.80 Gbps per flow. The major performance opportunity is therefore in the software transfer pipeline rather than the physical link.",
        "These tests did not perform tokenization, forward passes, backpropagation, optimizer updates, or GPU training. They validate the server exchange path, not end-to-end LLM training capacity.",
    ):
        add_bullet(document, text, bullet_num_id)

    add_heading(document, "What Was Tested - and What Was Not", 1)
    add_body(
        document,
        "The test model was a real PyTorch state dictionary with the exact requested number of BF16 or FP16 elements. It was split into bounded tensor shards so FLARE handled the same total payload size as the target model class. The tensors were synthetic zeros rather than downloaded Hugging Face weights."
    )
    add_figure(
        document,
        workflow_path,
        6.35,
        "Five-step diagram showing the server creating exact-size tensors, sending a global model, clients checking and changing one value, clients returning the full model state, and the server aggregating and checkpointing.",
        "Figure 1. The measured federated data path. The client operation intentionally replaces real training with a deterministic correctness marker.",
    )

    add_heading(document, "What the server performed", 2)
    for text in (
        "Created and retained the exact-size global model state.",
        "Distributed that model concurrently through production NVFLARE services.",
        "Received a complete returned model state from every participating client.",
        "Wrote incoming tensor data to disk-backed offload files and materialized tensors as aggregation required them.",
        "Performed weighted FedAvg aggregation and persisted a new global checkpoint each round.",
        "Captured process, cgroup, disk, network-phase, pressure, OOM, and correctness evidence.",
    ):
        add_bullet(document, text, bullet_num_id)

    add_heading(document, "What each client performed", 2)
    for text in (
        "Received the full model state and verified a known sentinel value from the previous global round.",
        "Changed one deterministic scalar value according to the client identity.",
        "Returned the otherwise unchanged full model state to the server.",
    ):
        add_bullet(document, text, bullet_num_id)

    add_heading(document, "What was intentionally skipped", 2)
    for text in (
        "No dataset, tokenizer, prompts, or CIFAR-10 images were loaded.",
        "No transformer forward pass, loss calculation, gradient, backward pass, or optimizer step was performed.",
        "No activations, optimizer states, or training checkpoints were allocated on the clients.",
        "No completed capacity run used a GPU, and no published model weights were downloaded.",
    ):
        add_bullet(document, text, bullet_num_id)
    add_callout(
        document,
        "Why this is still useful",
        "The expensive server operations - sending the entire model, receiving full updates, aggregating them, offloading tensors, and writing checkpoints - were exercised with real byte counts. The test deliberately isolates those operations from client-side training variability.",
    )

    document.add_page_break()
    add_heading(document, "Measured Results", 1)
    add_body(
        document,
        "The table below selects the representative distributed qualification result for each model size. It does not include setup retries, validation-only exports, or repeated smoke diagnostics. Lower tiers through 7.61B used two rounds; the 10B-and-larger scale points shown here used one round so the largest tiers could be compared safely."
    )
    scale_rows = [
        ["135M", "BF16", "0.25", "2 / 2", "1.91", "1.74", "4.8 min", "Pass"],
        ["0.49B", "BF16", "0.91", "2 / 2", "4.83", "2.79", "16.3 min", "Pass"],
        ["1.54B", "BF16", "2.87", "2 / 2", "12.39", "4.89", "3.0 min", "Pass"],
        ["3.09B", "BF16", "5.76", "2 / 2", "24.00", "8.61", "5.4 min", "Pass"],
        ["7.61B", "BF16", "14.17", "2 / 2", "57.97", "17.79", "12.8 min", "Pass"],
        ["10.0B", "BF16", "18.63", "1 / 2", "57.38", "22.00", "8.3 min", "Pass"],
        ["12.0B", "BF16", "22.35", "1 / 2", "68.65", "25.91", "10.4 min", "Pass"],
        ["14.7B", "BF16", "27.38", "1 / 2", "83.94", "31.19", "13.1 min", "Pass"],
        ["32.5B", "BF16", "60.54", "1 / 2", "184.51", "66.44", "30.3 min", "Pass"],
    ]
    add_table(
        document,
        ["Model size", "Type", "Payload\nGiB", "Rounds /\nclients", "Server peak\nGiB", "Largest client\nGiB", "Harness\nduration", "Outcome"],
        scale_rows,
        [1000, 700, 850, 1050, 1100, 1250, 1250, 2160],
        alignments=["left", "center", "right", "center", "right", "right", "center", "center"],
        font_size=8.6,
        status_column=7,
    )
    add_body(
        document,
        "GB versus GiB: model vendors use decimal billions, while this report expresses memory in binary GiB. For example, 65.0 GB and 60.54 GiB are the same number of bytes."
    )
    add_figure(
        document,
        scale_path,
        6.15,
        "Line chart for 10B, 12B, 14.7B, and 32.5B one-round two-client runs. Server process peak rises from 57.38 to 184.51 GiB, client process peak from 22.00 to 66.44 GiB, and raw payload from 18.63 to 60.54 GiB.",
        "Figure 2. Comparable one-round results show an approximately linear relationship between raw payload and process memory.",
    )

    add_heading(document, "Configuration comparisons", 2)
    comparison_rows = [
        ["14.7B BF16", "1 round, 2 clients", "83.94", "13.1 min", "Pass", "Large-model baseline"],
        ["14.7B BF16", "2 rounds, 2 clients", "111.08", "26.6 min", "Pass", "Both rounds completed"],
        ["14.7B BF16", "3 rounds, 2 clients", "110.97", "40.6 min", "Completed + warning", "RSS increased once, then plateaued"],
        ["14.7B FP16", "1 round, 2 clients", "83.44", "13.0 min", "Pass", "Equivalent capacity to BF16"],
        ["14.7B BF16", "8 MiB tensor requests", "83.37", "14.0 min", "Pass", "Slower than 2 MiB baseline"],
        ["32.5B BF16", "1 round, 2 clients", "184.51", "30.3 min", "Pass", "Fits 256/128 GiB class for one round"],
        ["32.5B BF16", "2 rounds, 2 clients, 235 GiB cap", "226.03", "63.0 min", "Stall; induced OOM", "Round 1 returned updates but could not aggregate"],
        ["32.5B BF16", "2 rounds, 3 clients, natural limit", "245.39", "58.7 min", "Pass", "All six client-round updates passed"],
        ["32.5B BF16", "3 rounds, 3 clients, natural limit", "245.46", "114.0 min", "Completed + warning", "All nine updates passed; final RSS plateaued"],
    ]
    add_table(
        document,
        ["Model", "Variation", "Server peak\nGiB", "Harness\nduration", "Outcome", "Meaning"],
        comparison_rows,
        [1150, 2080, 880, 970, 1290, 2990],
        alignments=["left", "left", "right", "center", "center", "left"],
        font_size=8.3,
        status_column=4,
    )
    add_body(
        document,
        "The 114-minute three-round duration includes status polling and artifact collection. The server's active federated workload lasted approximately 86 minutes. The formal warning was caused by a first-to-last round-end RSS rule; the final transition was approximately 0.02%, which is consistent with a one-time warm-up or retained allocation rather than continuing linear growth."
    )

    document.add_page_break()
    add_heading(document, "How Server Memory Works", 1)
    add_body(
        document,
        "A single 'memory used' number is not enough to size this workload. Linux reports several overlapping views. The most important distinction is between active application memory and file cache."
    )
    memory_definition_rows = [
        ["Raw model payload", "The tensor bytes for one model state.", "A 32.5B BF16 payload is 60.54 GiB."],
        ["Process RSS", "RAM currently associated with the server process tree.", "Primary measure of the live application working set."],
        ["Anonymous memory", "Tensors, heaps, serialization buffers, and allocator pages that are not simply a disk-file cache.", "Linux generally cannot discard it; the application must release it."],
        ["File cache", "Copies of recently read or written file data kept in RAM.", "Usually reclaimable when applications need more memory."],
        ["Kernel memory", "Filesystem, page-table, networking, and other operating-system structures.", "Small but necessary additional capacity."],
        ["Cgroup total", "Anonymous + file cache + kernel memory charged to the workload.", "Can look nearly full even when much of it is reclaimable."],
        ["MemAvailable", "Linux's estimate of memory that can be provided without severe disruption.", "Includes reclaimable cache; it is not a separate pool to add."],
    ]
    add_table(
        document,
        ["Term", "Plain-English meaning", "How to use it"],
        memory_definition_rows,
        [1550, 4050, 3760],
        alignments=["left", "left", "left"],
        font_size=8.8,
    )

    add_heading(document, "The largest measured memory point", 2)
    add_body(
        document,
        "The highest sampled cgroup total occurred late in the final round of the 32.5B three-client run. It was not 493.91 GiB of live tensors. More than half was file cache created by disk offload and checkpoint activity."
    )
    add_figure(
        document,
        memory_path,
        6.25,
        "Stacked bar showing the 493.91 GiB sampled cgroup peak: 224.70 GiB anonymous memory, 261.57 GiB file cache, and 7.63 GiB kernel memory. It notes 225.16 GiB process RSS and 268.20 GiB MemAvailable at the same instant.",
        "Figure 3. The cgroup peak was mostly a combination of active application memory and reclaimable file cache.",
    )
    add_callout(
        document,
        "Important interpretation",
        "493.91 GiB cgroup usage and 268.20 GiB MemAvailable are not contradictory. Linux charged the file cache to the workload while also recognizing that much of the same cache could be reclaimed.",
        fill=LIGHT_GOLD,
        color=GOLD,
    )

    add_heading(document, "Why more clients did not create full server copies", 2)
    add_body(
        document,
        "The server owns one global model. F3 is designed to cache each serialized tensor item, share the cached bytes among receivers, and remove the item after all receivers acknowledge it. Simultaneous first requests can transiently duplicate serialization work, but they do not add one permanent 60.54 GiB destination copy per client. More clients mainly add stream state, socket buffers, network traffic, and the possibility that a slower receiver retains shared bytes longer."
    )
    add_figure(
        document,
        client_path,
        5.65,
        "Bar chart showing comparable first-round server RSS of 180.33 GiB with one client, 184.51 GiB with two clients, and 185.09 GiB with three clients.",
        "Figure 4. From one to three clients, comparable first-round server RSS increased only 4.76 GiB, or 2.64%.",
    )
    add_body(
        document,
        "Each client still needs its own machine memory. At 32.5B, each client used approximately 66-67 GiB. Each client also returns a logically distinct full update, but disk offload writes those updates to files and aggregation uses one accumulator, updating it in place for later clients. More clients therefore increase network bytes, offload traffic, page-cache churn, and elapsed time more than permanent server RSS."
    )

    document.add_page_break()
    add_heading(document, "Understanding the Earlier Two-Client OOM", 1)
    add_callout(
        document,
        "Correction to the early explanation",
        "The server did not OOM because it held one full model copy for each client. The failure was a cgroup-constrained active-memory boundary encountered while preparing the next aggregation step.",
        fill=LIGHT_RED,
        color=RED,
    )
    add_body(
        document,
        "The earlier 32.5B two-client, two-round run used a 251 GiB-class server with a 235 GiB hard cgroup ceiling. Round zero completed. In round one, both clients received the model, returned valid full updates, and passed their correctness checks. The server then remained near 226 GiB without reaching aggregation."
    )
    for text in (
        "Anonymous memory was approximately 225.85 GiB, while ordinary file cache had already been reclaimed to approximately zero.",
        "Memory pressure averaged roughly 80% over ten seconds, and the server made no forward progress for more than seven minutes.",
        "The original 235 GiB limit did not naturally OOM during the observation window.",
        "To preserve lease time and obtain a complete failure artifact, the operator lowered the hard limit below current live usage. The final 225 GiB setting intentionally triggered the recorded cgroup OOM.",
    ):
        add_bullet(document, text, bullet_num_id)
    add_body(
        document,
        "Without the operator adjustment, the run might have remained stalled or eventually hit the original hard limit; the data cannot prove the exact counterfactual. However, the later unconstrained three-client run reached approximately 245.5 GiB process RSS. This shows that the old 235 GiB ceiling was below the active working set observed when the workflow completed."
    )

    add_heading(document, "Approximate capacity bands for 32.5B multi-round testing", 2)
    capacity_rows = [
        ["235 GiB cgroup", "Below demonstrated active requirement", "Observed stall; induced OOM after limit reduction", "Failure-boundary evidence only"],
        ["256 GiB class", "Almost no burst/cache headroom", "Extremely risky for multi-round work", "Do not use as a completion target"],
        ["About 288 GiB", "Roughly 35 GiB above active + kernel estimate", "Plausible bare-minimum pressure test", "May stall or vary substantially"],
        ["About 320 GiB", "Roughly 65-70 GiB cache/headroom", "Likely workable but reclaim-sensitive", "Useful soft-pressure point"],
        ["384 GiB class", "Roughly 130 GiB cache/headroom", "Recommended balance", "Expected to be materially smoother"],
        ["512 GiB class", "Approximately 250 GiB cache/headroom", "Measured low-risk baseline", "Completed three clients / three rounds"],
    ]
    add_table(
        document,
        ["Server budget", "Approximate interpretation", "Expected behavior", "Recommended use"],
        capacity_rows,
        [1430, 2580, 2760, 2590],
        alignments=["left", "left", "left", "left"],
        font_size=8.6,
    )
    document.add_page_break()
    add_body(
        document,
        "These bands are engineering projections, not yet a measured slowdown curve. The quickest calibration is to keep a 512 GiB server's hard limit natural and test soft memory.high values at 384 GiB and 320 GiB. That allows Linux to reclaim and throttle without manufacturing another OOM."
    )

    add_heading(document, "Controlled failure evidence", 2)
    add_body(
        document,
        "Separate 10B controls verified that the harness distinguishes slow memory pressure from a true OOM. A soft-pressure lane accumulated approximately 281 seconds with all cgroup tasks stalled but recorded no OOM; a 42 GiB hard-limit lane produced the expected cgroup OOM after valid client returns. This matters operationally because a job can stop progressing long before the kernel kills it."
    )

    add_heading(document, "Performance Findings", 1)
    add_body(
        document,
        "Network measurements show that the machines and switch were not the primary transfer limit. Three concurrent iperf3 streams filled almost all of the available 25 Gbps-class link, while NVFLARE used a small fraction of that capacity."
    )
    performance_rows = [
        ["Concurrent iperf3, clients to server", "23.53 Gbps aggregate", "Physical path is capable of near-line-rate traffic"],
        ["Concurrent iperf3, server to clients", "23.65 Gbps aggregate", "Fan-out path is also capable of near-line-rate traffic"],
        ["NVFLARE model flow", "Approximately 0.75-0.80 Gbps per flow", "Application pipeline, not NIC capacity, dominates elapsed time"],
        ["Best F3 buffer profile", "4 MiB chunk / 128 MiB window / 32 MiB ACK", "Only 8% faster than default"],
        ["Larger tensor request batch", "2 MiB to 768 MiB", "No material improvement"],
    ]
    add_table(
        document,
        ["Measurement", "Observed result", "Conclusion"],
        performance_rows,
        [3000, 2450, 3910],
        alignments=["left", "center", "left"],
        font_size=9,
    )
    add_heading(document, "Most likely bottleneck", 2)
    add_body(
        document,
        "The strongest working diagnosis is a bulk-tensor software-pipeline efficiency problem. Whole tensors are serialized into safetensors bytes; the receiver requests and consumes batches synchronously; incoming items are written to disk; and gRPC/protobuf introduces additional copying. Production, transport, and consumption are not overlapped efficiently enough to approach the measured network capacity."
    )
    for text in (
        "A larger F3 window was not the answer: the default window already exceeded the measured bandwidth-delay product.",
        "A 4 MiB F3 profile was the best tested setting, but the gain was only 8% and larger settings regressed.",
        "Increasing tensor request batching by 384 times did not improve the critical path.",
        "The next meaningful work is a gRPC+TLS versus TCP+TLS comparison and bounded pipelining/single-flight serialization, not another generic buffer sweep.",
    ):
        add_bullet(document, text, bullet_num_id)

    add_heading(document, "Conclusions and Recommended Next Steps", 1)
    add_numbered(
        document,
        "Measure the 32B memory curve at soft thresholds of 384 GiB and 320 GiB while leaving the hard limit natural. Record phase duration, PSI pressure, reclaim events, disk I/O, page faults, anonymous memory, and file cache.",
        decimal_num_id,
    )
    add_numbered(
        document,
        "Run a controlled gRPC+TLS versus TCP+TLS transfer comparison with the 1.5B fast driver, then confirm any improvement at 14B and 32B.",
        decimal_num_id,
    )
    add_numbered(
        document,
        "Add one real-model confirmation at 14.7B: load pinned published weights, use one tiny fixed batch, and update only one selected layer or adapter for one step while returning the full state dictionary.",
        decimal_num_id,
    )
    add_numbered(
        document,
        "Defer a 72B capacity run until transfer throughput improves materially. Current planning ranges are approximately 390-450 GiB server process RSS for one round and 520-650 GiB for two rounds, with 256 GiB clients.",
        decimal_num_id,
    )
    add_callout(
        document,
        "Recommended machine starting points",
        "For another 32B multi-round qualification, use a 384 GiB server as the balanced target or a 512 GiB server as the measured low-risk baseline, with 128 GiB per client. For an eventual 72B one-round gate, use a 768 GiB boundary server or 1 TiB low-risk server and 256 GiB per client.",
        fill=LIGHT_GREEN,
        color=GREEN,
    )

    document.add_page_break()
    add_heading(document, "Appendix A - Evidence Boundary and Selected Runs", 1)
    add_body(
        document,
        "Every harness invocation created an immutable directory containing the exact scenario, source revision, events, process and cgroup samples, service logs, result status, and machine-readable report. The following run IDs anchor the conclusions in this document."
    )
    run_rows = [
        ["Distributed scaling", "135M through 12B", "20260715T185906Z... through 20260715T213902Z...", "Passing production qualification curve"],
        ["14.7B baseline", "BF16, 1 round, 2 clients", "20260717T161123Z-qwen25-14b-shape-1r-ac219ce8", "Pass"],
        ["14.7B repeatability", "BF16, 2 and 3 rounds", "20260717T163651Z... / 20260717T203906Z...", "Completed; three-round growth warning"],
        ["14.7B dtype/chunk", "FP16 and 8 MiB comparison", "20260717T185252Z... / 20260717T192816Z...", "Both completed"],
        ["32.5B baseline", "BF16, 1 round, 2 clients", "20260717T195343Z-qwen25-32b-shape-1r-91de0834", "Pass"],
        ["32.5B pressure", "BF16, 2 rounds, 2 clients", "20260717T213434Z-qwen25-32b-shape-39c994bf", "Soft stall; induced OOM capture"],
        ["32.5B three-client", "BF16, 2 rounds, 3 clients", "20260720T205721Z-qwen25-32b-3client-f3-4m-batch2m-2r-d6e01997", "Pass"],
        ["32.5B final", "BF16, 3 rounds, 3 clients", "20260720T224359Z-qwen25-32b-3client-f3-4m-batch2m-3r-bce5fd11", "Completed; conservative growth warning"],
    ]
    add_table(
        document,
        ["Evidence group", "Configuration", "Representative run ID", "Result"],
        run_rows,
        [1500, 2350, 3800, 1710],
        alignments=["left", "left", "left", "left"],
        font_size=7.9,
        status_column=3,
    )

    add_heading(document, "Limits of interpretation", 2)
    for text in (
        "Synthetic tensors reproduce payload size and the server exchange path, but not the tensor topology, library behavior, or GPU memory of a real transformer training job.",
        "Client training memory cannot be inferred from these runs because activations, gradients, and optimizer state were intentionally absent.",
        "The memory sizing bands are based on the measured environment and may shift with operating system, allocator, storage, transport driver, model tensor topology, or NVFLARE version.",
        "File-cache peaks are timing-sensitive. Category maxima should not be added unless they occurred at the same sample.",
        "The exact minimum smooth server size still requires the proposed soft-threshold sweep.",
    ):
        add_bullet(document, text, bullet_num_id)

    document.add_page_break()
    add_heading(document, "Plain-language glossary", 2)
    glossary_rows = [
        ["Federated round", "One cycle in which the server sends a global model, clients return updates, and the server combines them."],
        ["BF16 / FP16", "Two-byte numeric formats. Both require the same raw payload size in these tests."],
        ["Sentinel", "A known value used to prove that each client received the expected model and that aggregation incorporated the update."],
        ["Disk tensor offload", "Writing incoming tensor bytes to files and loading pieces only when aggregation needs them."],
        ["RSS", "Resident set size: physical RAM currently associated with a process."],
        ["Cgroup", "A Linux accounting and control boundary that can measure or limit a group of processes."],
        ["OOM", "Out of memory. In these tests, a cgroup OOM means the workload crossed its configured hard ceiling; it does not necessarily mean the entire machine exhausted every byte."],
        ["PSI", "Linux pressure-stall information, used to measure how long tasks wait because memory is under pressure."],
    ]
    add_table(
        document,
        ["Term", "Meaning"],
        glossary_rows,
        [1800, 7560],
        alignments=["left", "left"],
        font_size=8.8,
    )

    add_body(
        document,
        "Report prepared from retained NVFLARE stress-test artifacts through July 21, 2026. Values are rounded for readability; machine-readable artifacts retain byte-level measurements.",
    )

    for section in document.sections:
        setup_page(section)

    document.core_properties.title = "NVFLARE Large-Model Server Stress Testing"
    document.core_properties.subject = "External technical findings through 32.5B, three clients, and three rounds"
    document.core_properties.author = "NVFLARE Stress Test Project"
    document.core_properties.keywords = "NVFLARE, LLM, federated learning, memory, stress testing"
    document.core_properties.comments = "Prepared from retained immutable run artifacts."
    document.save(OUTPUT_PATH)
    return OUTPUT_PATH


if __name__ == "__main__":
    print(build_report())
